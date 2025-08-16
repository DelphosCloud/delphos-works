from flask import Flask, request, jsonify, send_file
import json
import os
import uuid
import boto3
from docx import Document
from datetime import datetime
import io
import tempfile
from botocore.exceptions import ClientError
import base64

app = Flask(__name__)

@app.route('/api/generate', methods=['POST'])
def generate_document():
    """
    Generate document from template and JSON data
    """
    try:
        # Get JSON data from request
        data = request.get_json()
        
        if not data:
            return jsonify({
                "Title": "Error",
                "Message": "JSON data is required"
            }), 400
        
        # Validate required templateId
        template_id = data.get('templateId')
        if not template_id:
            return jsonify({
                "Title": "Error",
                "Message": "templateId is required"
            }), 400
        
        # Initialize Spaces client
        spaces_client = boto3.client(
            's3',
            endpoint_url=f"https://{os.environ['SPACES_ENDPOINT']}",
            aws_access_key_id=os.environ['SPACES_KEY'],
            aws_secret_access_key=os.environ['SPACES_SECRET']
        )
        
        bucket_name = os.environ['SPACES_BUCKET']
        
        # Download template from Spaces
        try:
            template_key = f"templates/{template_id}"
            template_response = spaces_client.get_object(Bucket=bucket_name, Key=template_key)
            template_content = template_response['Body'].read()
        except Exception as e:
            return jsonify({
                "Title": "Error",
                "Message": "The Template could not be found."
            }), 404
        
        # Load the document from memory
        template_doc = Document(io.BytesIO(template_content))
        
        # Replace placeholders in paragraphs
        for paragraph in template_doc.paragraphs:
            replace_placeholders_in_text(paragraph, data)
        
        # Handle table processing for deliverables
        process_tables(template_doc, data)
        
        # Generate unique filename for output
        output_filename = str(uuid.uuid4())
        
        # Save document to memory buffer
        doc_buffer = io.BytesIO()
        template_doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Upload Word document to Spaces
        word_key = f"generated/{output_filename}.docx"
        spaces_client.put_object(
            Bucket=bucket_name,
            Key=word_key,
            Body=doc_buffer.getvalue(),
            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        result = {
            "fileWordDoc": f"{output_filename}.docx",
            "filePdfDoc": "",
            "timeStamp": datetime.now().isoformat()
        }
        
        # Generate PDF if requested (placeholder for now)
        if data.get('pdf', False):
            # PDF generation would go here
            result["filePdfDoc"] = ""
        
        return jsonify(result), 200
        
    except Exception as e:
        return jsonify({
            "Title": "Error",
            "Message": f"Internal server error: {str(e)}"
        }), 500

@app.route('/api/download', methods=['GET'])
def download_file():
    """
    Download generated files
    """
    try:
        # Get filename from query parameters
        filename = request.args.get('file')
        
        if not filename:
            return jsonify({
                "Title": "Error",
                "Message": "Filename parameter is required"
            }), 400
        
        # Initialize Spaces client
        spaces_client = boto3.client(
            's3',
            endpoint_url=f"https://{os.environ['SPACES_ENDPOINT']}",
            aws_access_key_id=os.environ['SPACES_KEY'],
            aws_secret_access_key=os.environ['SPACES_SECRET']
        )
        
        bucket_name = os.environ['SPACES_BUCKET']
        file_key = f"generated/{filename}"
        
        try:
            # Get the file from Spaces
            file_response = spaces_client.get_object(Bucket=bucket_name, Key=file_key)
            file_content = file_response['Body'].read()
            
            # Create a temporary file to serve
            temp_file = io.BytesIO(file_content)
            
            # Determine content type
            mimetype = 'application/octet-stream'
            if filename.endswith('.docx'):
                mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif filename.endswith('.pdf'):
                mimetype = 'application/pdf'
            
            return send_file(
                temp_file,
                as_attachment=True,
                download_name=filename,
                mimetype=mimetype
            )
            
        except ClientError as e:
            if e.response['Error']['Code'] == 'NoSuchKey':
                return jsonify({
                    "Title": "Error",
                    "Message": "File not found"
                }), 404
            else:
                raise e
                
    except Exception as e:
        return jsonify({
            "Title": "Error",
            "Message": f"Internal server error: {str(e)}"
        }), 500

def replace_placeholders_in_text(paragraph, data):
    """Replace placeholders in paragraph text"""
    full_text = paragraph.text
    
    # Replace simple placeholders
    for key, value in data.items():
        if key != 'deliverables' and isinstance(value, (str, int, float)):
            placeholder = f"{{{{{key}}}}}"
            full_text = full_text.replace(placeholder, str(value))
    
    # Clear and rebuild paragraph with replaced text
    if full_text != paragraph.text:
        paragraph.clear()
        paragraph.add_run(full_text)

def process_tables(doc, data):
    """Process tables and handle deliverables repetition"""
    deliverables = data.get('deliverables', [])
    
    for table in doc.tables:
        rows_to_remove = []
        rows_to_add = []
        
        for i, row in enumerate(table.rows):
            row_text = ' '.join(cell.text for cell in row.cells)
            
            # Check if this row contains the repeat marker
            if '{{!REPEATROW}}' in row_text:
                # Store the template row
                template_row = row
                rows_to_remove.append(i)
                
                # Create new rows for each deliverable
                for deliverable in deliverables:
                    new_row_data = []
                    for cell in template_row.cells:
                        cell_text = cell.text
                        # Remove the repeat marker
                        cell_text = cell_text.replace('{{!REPEATROW}}', '')
                        
                        # Replace deliverable placeholders
                        for key, value in deliverable.items():
                            placeholder = f"{{{{{key}}}}}"
                            cell_text = cell_text.replace(placeholder, str(value))
                        
                        new_row_data.append(cell_text)
                    
                    rows_to_add.append((i, new_row_data))
        
        # Remove template rows (in reverse order to maintain indices)
        for row_index in reversed(rows_to_remove):
            table._element.remove(table.rows[row_index]._element)
        
        # Add new rows
        for original_index, row_data in rows_to_add:
            new_row = table.add_row()
            for j, cell_text in enumerate(row_data):
                if j < len(new_row.cells):
                    new_row.cells[j].text = cell_text

@app.route('/', methods=['GET'])
def health_check():
    """Basic health check endpoint"""
    return jsonify({"status": "healthy", "service": "delphos-works document generator"})

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080)

@app.route('/debug/env', methods=['GET'])
def debug_env():
    """Debug endpoint to check environment variables"""
    return jsonify({
        "SPACES_ENDPOINT": os.environ.get('SPACES_ENDPOINT', 'NOT_SET'),
        "SPACES_BUCKET": os.environ.get('SPACES_BUCKET', 'NOT_SET'),
        "SPACES_KEY": os.environ.get('SPACES_KEY', 'NOT_SET')[:10] + "..." if os.environ.get('SPACES_KEY') else 'NOT_SET',
        "SPACES_SECRET": "SET" if os.environ.get('SPACES_SECRET') else 'NOT_SET'
    })
