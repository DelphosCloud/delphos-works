import json
import os
import uuid
import boto3
from docx import Document
from docx.shared import Inches
from datetime import datetime
import io
import tempfile

def main(event, context):
    """
    Main handler for document generation
    Receives JSON, processes template, returns generated file info
    """
    try:
        # Parse the incoming JSON
        if isinstance(event.get('body'), str):
            body = json.loads(event['body'])
        else:
            body = event.get('body', {})
        
        # Validate required templateId
        template_id = body.get('templateId')
        if not template_id:
            return {
                'statusCode': 400,
                'body': json.dumps({
                    "Title": "Error",
                    "Message": "templateId is required"
                })
            }
        
        # Initialize Spaces client
        spaces_client = boto3.client(
            's3',
            endpoint_url=f"https://{os.environ['SPACES_ENDPOINT']}",
            aws_access_key_id=os.environ['SPACES_KEY'],
            aws_secret_access_key=os.environ['SPACES_SECRET']
        )
        
        bucket_name = os.environ['SPACES_BUCKET']
        
        # Download template from Spaces
        try:
            template_key = f"templates/{template_id}"
            template_response = spaces_client.get_object(Bucket=bucket_name, Key=template_key)
            template_content = template_response['Body'].read()
        except Exception as e:
            return {
                'statusCode': 404,
                'body': json.dumps({
                    "Title": "Error",
                    "Message": "The Template could not be found."
                })
            }
        
        # Load the document from memory
        template_doc = Document(io.BytesIO(template_content))
        
        # Replace placeholders in paragraphs
        for paragraph in template_doc.paragraphs:
            replace_placeholders_in_text(paragraph, body)
        
        # Handle table processing for deliverables
        process_tables(template_doc, body)
        
        # Generate unique filename for output
        output_filename = str(uuid.uuid4())
        
        # Save document to memory buffer
        doc_buffer = io.BytesIO()
        template_doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Upload Word document to Spaces
        word_key = f"generated/{output_filename}.docx"
        spaces_client.put_object(
            Bucket=bucket_name,
            Key=word_key,
            Body=doc_buffer.getvalue(),
            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        result = {
            "fileWordDoc": f"{output_filename}.docx",
            "filePdfDoc": "",
            "timeStamp": datetime.now().isoformat()
        }
        
        # Generate PDF if requested
        if body.get('pdf', False):
            try:
                # For now, we'll skip PDF generation as it requires additional libraries
                # You can add PDF conversion here later if needed
                result["filePdfDoc"] = ""
            except Exception as pdf_error:
                # Continue without PDF if conversion fails
                result["filePdfDoc"] = ""
        
        return {
            'statusCode': 200,
            'body': json.dumps(result),
            'headers': {
                'Content-Type': 'application/json'
            }
        }
        
    except Exception as e:
        return {
            'statusCode': 500,
            'body': json.dumps({
                "Title": "Error",
                "Message": f"Internal server error: {str(e)}"
            })
        }

def replace_placeholders_in_text(paragraph, data):
    """Replace placeholders in paragraph text"""
    full_text = paragraph.text
    
    # Replace simple placeholders
    for key, value in data.items():
        if key != 'deliverables' and isinstance(value, (str, int, float)):
            placeholder = f"{{{{{key}}}}}"
            full_text = full_text.replace(placeholder, str(value))
    
    # Clear and rebuild paragraph with replaced text
    if full_text != paragraph.text:
        paragraph.clear()
        paragraph.add_run(full_text)

def process_tables(doc, data):
    """Process tables and handle deliverables repetition"""
    deliverables = data.get('deliverables', [])
    
    for table in doc.tables:
        rows_to_remove = []
        rows_to_add = []
        
        for i, row in enumerate(table.rows):
            row_text = ' '.join(cell.text for cell in row.cells)
            
            # Check if this row contains the repeat marker
            if '{{!REPEATROW}}' in row_text:
                # Store the template row
                template_row = row
                rows_to_remove.append(i)
                
                # Create new rows for each deliverable
                for deliverable in deliverables:
                    new_row_data = []
                    for cell in template_row.cells:
                        cell_text = cell.text
                        # Remove the repeat marker
                        cell_text = cell_text.replace('{{!REPEATROW}}', '')
                        
                        # Replace deliverable placeholders
                        for key, value in deliverable.items():
                            placeholder = f"{{{{{key}}}}}"
                            cell_text = cell_text.replace(placeholder, str(value))
                        
                        new_row_data.append(cell_text)
                    
                    rows_to_add.append((i, new_row_data))
        
        # Remove template rows (in reverse order to maintain indices)
        for row_index in reversed(rows_to_remove):
            table._element.remove(table.rows[row_index]._element)
        
        # Add new rows
        for original_index, row_data in rows_to_add:
            new_row = table.add_row()
            for j, cell_text in enumerate(row_data):
                if j < len(new_row.cells):
                    new_row.cells[j].text = cell_text